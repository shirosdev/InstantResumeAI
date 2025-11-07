# backend/app/routes/resume.py
# --- FULL RECTIFIED FILE ---

from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import uuid
import re
from app.models.subscription import UserSubscription, SubscriptionPlan 
from datetime import datetime
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.services.auth_service import AuthService

import docx
from app import db
from app.models.user import User
from app.models.activity import ActivityLog
# Import the new advanced processor instead of the old one
from app.services.comprehensive_resume_processor import IntegratedResumeService
from app.services.enhancement_validator import EnhancementValidator
from app.models.resume import ResumeEnhancement

resume_bp = Blueprint('resume', __name__)

# Configure upload settings
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'uploads')
ALLOWED_EXTENSIONS = {'docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_INSTRUCTION_LENGTH = 2000  # Increased for detailed instructions

# Ensure upload directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'resumes'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'enhanced'), exist_ok=True)

def print_progress(step, total_steps, description):
    """A simple callback function to print progress to the console."""
    percentage = int((step / total_steps) * 100)
    print(f"[PROGRESS {percentage}%] Step {step}/{total_steps}: {description}")

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file_size(file):
    """Validate file size is within limits"""
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)
    return file_size <= MAX_FILE_SIZE

def validate_user_instructions(instructions):
    """Enhanced validation for user instructions"""
    if not instructions:
        return True, ""
    
    if len(instructions) > MAX_INSTRUCTION_LENGTH:
        return False, f"Instructions must be less than {MAX_INSTRUCTION_LENGTH} characters"
    
    harmful_patterns = [
        (r'\b(?:create|invent|fabricate)\s+(?:fake|false|fictitious)\s+\w+', 
         "Instructions must be based on truthful information"),
        (r'\b(?:illegal|unethical|fraudulent)\b', 
         "Instructions must be professional and ethical")
    ]
    
    for pattern, error_msg in harmful_patterns:
        if re.search(pattern, instructions, re.IGNORECASE):
            return False, error_msg
    
    return True, ""

def get_text_from_docx(filepath):
    """Extract all text from a DOCX file"""
    try:
        doc = docx.Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
        return ""


@resume_bp.route('/enhance', methods=['POST'])
@jwt_required()
def enhance_resume():
    """Main enhancement endpoint with advanced processing"""
    user_id_from_token = int(get_jwt_identity())
    user = User.query.get(user_id_from_token)
    
    if not user or not user.is_active:
        return jsonify({'message': 'User not found or is inactive'}), 401

    # --- CREDIT CHECK LOGIC ---
    try:
        db.session.expire_all()
        user_subscription = UserSubscription.query.filter_by(user_id=user_id_from_token).first()
        if not user_subscription:
            return jsonify({'message': 'User has no active subscription plan.'}), 403

        plan = SubscriptionPlan.query.get(user_subscription.plan_id)
        if not plan:
             return jsonify({'message': 'Could not verify subscription plan.'}), 500

        enhancement_count = ResumeEnhancement.query.filter_by(user_id=user_id_from_token).count()
        
        total_available = 0
        if plan.resume_limit is not None:
            total_available = plan.resume_limit + (user_subscription.enhancement_credits or 0)
        else:
            total_available = float('inf') 

        print(f"[CREDIT CHECK] User {user_id_from_token}: Used={enhancement_count}, Available={total_available}")

        if enhancement_count >= total_available:
            return jsonify({
                'message': 'You have used all your free enhancements and credits. Please top-up to continue.',
                'limit_reached': True
            }), 403

    except Exception as e:
        db.session.rollback()
        print(f"Error in credit check: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({'message': 'Could not verify subscription status', 'error': str(e)}), 500
        
    # --- File Validation and Processing ---
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({'message': 'Resume file (.docx) and job description are required'}), 400
    
    resume_file = request.files['resume']
    job_description = request.form.get('job_description', '').strip()
    user_instructions = request.form.get('user_instructions', '').strip()
    
    if not resume_file.filename or not allowed_file(resume_file.filename):
        return jsonify({'message': 'Invalid file type. Please upload a .docx file'}), 400
    
    if not validate_file_size(resume_file):
        return jsonify({'message': 'File size exceeds 10MB limit'}), 400
    
    if len(job_description) < 50:
        return jsonify({'message': 'Job description must be at least 50 characters'}), 400
    
    instructions_valid, instructions_error = validate_user_instructions(user_instructions)
    if not instructions_valid:
        return jsonify({'message': instructions_error}), 400

    user_facing_filename = secure_filename(resume_file.filename)
    unique_id = str(uuid.uuid4())
    internal_filename = f"{unique_id}.docx"
    filepath = os.path.join(UPLOAD_FOLDER, 'resumes', internal_filename)
    
    try:
        resume_file.save(filepath)
    except Exception as e:
        return jsonify({'message': 'Failed to save uploaded file', 'error': str(e)}), 500

    # --- Main Enhancement and Database Update Logic ---
    try:
        enhancement_service = IntegratedResumeService()
        enhanced_filename = f"enhanced_{unique_id}.docx"
        enhanced_filepath = os.path.join(UPLOAD_FOLDER, 'enhanced', enhanced_filename)
        
        print(f"Starting enhancement for user {user_id_from_token}")
        print(f"Job Description Length: {len(job_description)}")
        print(f"User Instructions: {'Yes' if user_instructions else 'No'}")
        
        processing_result = enhancement_service.enhance_resume(
            original_file_path=filepath,
            job_description=job_description,
            user_instructions=user_instructions if user_instructions else None,
            output_path=enhanced_filepath,
            progress_callback=print_progress
        )
        
        if not processing_result.get('success'):
            if os.path.exists(filepath): os.remove(filepath)
            error_detail = processing_result.get('error', 'Unknown processing error')
            print(f"Enhancement failed: {error_detail}")
            return jsonify({'message': 'Failed to enhance resume', 'error': error_detail}), 500
        
        validator = EnhancementValidator()
        original_text = get_text_from_docx(filepath)
        enhanced_text = get_text_from_docx(enhanced_filepath)
        
        validation_results = validator.validate_enhancement(
            original_text=original_text,
            enhanced_text=enhanced_text,
            job_description=job_description
        )
        
        final_file_size = os.path.getsize(enhanced_filepath)
        
        new_enhancement = ResumeEnhancement(
            user_id=user_id_from_token,
            original_filename=user_facing_filename,
            enhanced_filename=enhanced_filename,
            file_path=os.path.join('enhanced', enhanced_filename),
            enhancement_status='completed',
            job_description_snippet=job_description[:500],
            enhancement_summary=user_instructions if user_instructions else None,
            created_at=datetime.utcnow(),
            completed_at=datetime.utcnow()
        )
        db.session.add(new_enhancement)

        activity_description = f'Enhanced resume: {user_facing_filename}'
        if user_instructions:
            activity_description += ' (with custom instructions)'
        
        new_activity_log = ActivityLog(user_id=user_id_from_token, action='resume_enhanced', description=activity_description)
        db.session.add(new_activity_log)
        
        if plan.resume_limit is not None and enhancement_count >= plan.resume_limit:
            print(f"User {user_id_from_token} is using a purchased credit. Attempting to decrement.")
            
            db.session.query(UserSubscription).filter(
                UserSubscription.user_id == user_id_from_token,
                UserSubscription.enhancement_credits > 0
            ).update({
                'enhancement_credits': UserSubscription.enhancement_credits - 1,
                'updated_at': datetime.utcnow()
            }, synchronize_session=False)
        
        db.session.commit()

        cleanup_old_enhancements(user_id_from_token, plan)

        response_data = {
            'message': 'Resume enhanced successfully',
            'enhanced_resume_id': unique_id,
            'file_size': final_file_size,
            'custom_instructions_applied': bool(user_instructions),
            'enhancements_made': processing_result.get('enhancements_made', 0),
            'quality_assessment': {
                'score': validation_results.get('overall_score', 0),
                'issues': validation_results.get('issues_found', [])
            }
        }
        
        print(f"Enhancement completed successfully for user {user_id_from_token}")
        return jsonify(response_data), 200

    except Exception as e:
        import traceback
        db.session.rollback()
        
        if os.path.exists(filepath): os.remove(filepath)
        if 'enhanced_filepath' in locals() and os.path.exists(enhanced_filepath): os.remove(enhanced_filepath)
        
        error_trace = traceback.format_exc()
        print(f"Critical error during enhancement: {str(e)}")
        print(error_trace)
        
        return jsonify({'message': 'An unexpected error occurred during enhancement', 'error': str(e)}), 500

def cleanup_old_enhancements(user_id, plan):
    """Clean up old enhancement records and files, but only for users with a limit."""
    if plan and plan.resume_limit is None:
        print(f"Skipping cleanup for user {user_id} on unlimited plan.")
        return

    try:
        all_enhancements = ResumeEnhancement.query.filter_by(user_id=user_id)\
            .order_by(ResumeEnhancement.created_at.asc()).all()
        
        if len(all_enhancements) > 5:
            records_to_delete = all_enhancements[:-5]
            
            for record in records_to_delete:
                enhanced_file = os.path.join(UPLOAD_FOLDER, 'enhanced', record.enhanced_filename)
                if os.path.exists(enhanced_file):
                    os.remove(enhanced_file)
                
                original_file_id = record.enhanced_filename.replace('enhanced_', '').replace('.docx', '')
                original_file = os.path.join(UPLOAD_FOLDER, 'resumes', f"{original_file_id}.docx")
                if os.path.exists(original_file):
                    os.remove(original_file)
                
                db.session.delete(record)
            
            db.session.commit()
            print(f"Cleaned up {len(records_to_delete)} old enhancement records for user {user_id}")
            
    except Exception as e:
        print(f"Error during cleanup: {str(e)}")
        db.session.rollback()

@resume_bp.route('/history', methods=['GET'])
@jwt_required()
def get_enhancement_history():
    """Get user's enhancement history"""
    try:
        current_user_id = get_jwt_identity()
        
        history_records = ResumeEnhancement.query.filter_by(user_id=current_user_id)\
            .order_by(ResumeEnhancement.created_at.desc())\
            .limit(5).all()
        
        history_list = []
        for record in history_records:
            record_dict = record.to_dict()
            enhanced_file = os.path.join(UPLOAD_FOLDER, 'enhanced', record.enhanced_filename)
            record_dict['file_available'] = os.path.exists(enhanced_file)
            history_list.append(record_dict)
        
        return jsonify(history_list), 200
        
    except Exception as e:
        print(f"Error fetching history: {str(e)}")
        return jsonify({'message': 'Failed to retrieve enhancement history'}), 500

# --- ADD 'OPTIONS' TO THE METHODS LIST ---
@resume_bp.route('/download/<resume_id>', methods=['GET', 'OPTIONS'])
@jwt_required()
def download_enhanced_resume(resume_id):
    """Download enhanced resume file"""
    
    if not re.match(r'^[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}$', resume_id):
        return jsonify({'message': 'Invalid resume ID format'}), 400
    
    enhanced_filename = f"enhanced_{resume_id}.docx"
    filepath = os.path.join(UPLOAD_FOLDER, 'enhanced', enhanced_filename)
    
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return jsonify({
            'message': 'File not found. The download link may have expired.'
        }), 404
    
    download_name = f'Enhanced_Resume_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    
    try:
        return send_file(
            filepath,
            as_attachment=True,
            download_name=download_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error sending file: {str(e)}")
        return jsonify({'message': 'Error downloading file'}), 500

@resume_bp.route('/validate', methods=['POST'])
@jwt_required()
def validate_enhancement_input():
    """Validate enhancement inputs before processing"""
    try:
        data = request.get_json()
        
        job_description = data.get('job_description', '').strip()
        if len(job_description) < 50:
            return jsonify({
                'valid': False,
                'message': 'Job description must be at least 50 characters'
            }), 400
        
        user_instructions = data.get('user_instructions', '').strip()
        if user_instructions:
            instructions_valid, instructions_error = validate_user_instructions(user_instructions)
            if not instructions_valid:
                return jsonify({
                    'valid': False,
                    'message': instructions_error
                }), 400
        
        return jsonify({
            'valid': True,
            'message': 'Inputs are valid'
        }), 200
        
    except Exception as e:
        return jsonify({
            'valid': False,
            'message': f'Validation error: {str(e)}'
        }), 400

@resume_bp.route('/log-disclaimer-agreement', methods=['POST'])
@jwt_required()
def log_disclaimer_agreement():
    """Logs the user's acceptance of the pre-download disclaimer."""
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        enhancement_id = data.get('enhancement_id')

        if not enhancement_id:
            return jsonify({'message': 'Enhancement ID is required'}), 400

        AuthService.log_activity(
            user_id=current_user_id,
            action='disclaimer_accepted',
            description=f'User agreed to the disclaimer for enhancement ID: {enhancement_id}',
            ip_address=request.remote_addr
        )
        
        return jsonify({'message': 'Agreement logged successfully'}), 200

    except Exception as e:
        return jsonify({'message': 'Failed to log agreement', 'error': str(e)}), 500