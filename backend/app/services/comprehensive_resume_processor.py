import os
import re
import json
import asyncio
import aiohttp
import logging
import copy
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import diff_match_patch as dmp_module
from typing import Dict, List, Tuple, Optional, Callable # Make sure Callable is imported
from dataclasses import dataclass
from enum import Enum
from docx.text.paragraph import Paragraph
from docx.table import Table
from thefuzz import fuzz

class SectionType(Enum):
    """Enum for different resume sections"""
    CONTACT = "contact"
    SUMMARY = "summary"
    EXPERIENCE = "experience"
    EDUCATION = "education"
    SKILLS = "skills"
    PROJECTS = "projects"
    CERTIFICATIONS = "certifications"
    ACHIEVEMENTS = "achievements"
    UNKNOWN = "unknown"

@dataclass
class ResumeSection:
    """Data class for resume sections"""
    type: SectionType
    title: str
    title_index: int
    content: List[str]
    paragraph_indices: List[int]

class IntegratedResumeService:

    # --- ADD PROGRESS STEPS ---
    PROGRESS_STEPS = [
        {"step": 1, "description": "Analyzing resume structure and sections"},
        {"step": 2, "description": "Identifying key information to preserve"},
        {"step": 3, "description": "Generating new content based on instructions (if any)"},
        {"step": 4, "description": "Preparing content chunks for AI enhancement"},
        {"step": 5, "description": "Enhancing resume content with AI (processing chunks)"},
        {"step": 6, "description": "Applying enhancements while preserving formatting"},
        {"step": 7, "description": "Integrating newly generated content"},
        {"step": 8, "description": "Finalizing document and saving"},
    ]
    TOTAL_STEPS = len(PROGRESS_STEPS)
    # --- END ADD ---

    def __init__(self):
        self.xml_illegal_char_re = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]')
        self.logger = logging.getLogger(__name__)
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        self.client = self._initialize_openai()
        self.dmp = dmp_module.diff_match_patch()
        self.protection_patterns = self._initialize_protection_patterns()
        self.section_patterns = self._initialize_section_patterns()

    def _initialize_openai(self):
        try:
            api_key = os.getenv('OPENAI_API_KEY')
            if not api_key:
                self.logger.error("OpenAI API key not found in environment variables.")
                return None
            self.logger.info("OpenAI client initialized successfully")
            return OpenAI(api_key=api_key)
        except Exception as e:
            self.logger.error(f"OpenAI initialization failed: {str(e)}")
            return None

    def _initialize_section_patterns(self):
        """Initialize patterns for identifying resume sections"""
        return {
            SectionType.EXPERIENCE: [r'(?i)^(professional\s+)?experience', r'(?i)^work\s+(history|experience)', r'(?i)^employment(\s+history)?'],
            SectionType.EDUCATION: [r'(?i)^education', r'(?i)^academic\s+(background|qualifications)'],
            SectionType.SKILLS: [r'(?i)^(technical\s+)?skills', r'(?i)^core\s+competencies', r'(?i)^expertise'],
            SectionType.PROJECTS: [r'(?i)^projects?', r'(?i)^(selected\s+)?projects'],
            SectionType.SUMMARY: [r'(?i)^(professional\s+)?summary', r'(?i)^objective', r'(?i)^profile'],
            SectionType.CERTIFICATIONS: [r'(?i)^certifications?', r'(?i)^licenses?\s*(and|&)?\s*certifications?']
        }

    def _initialize_protection_patterns(self):
        """Initializes regex patterns for identifying critical, non-negotiable text."""
        return {
            'dates': [
                r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\'?\d{2,4}\b',
                r'\b\d{4}\s*[-–—]\s*(?:\d{4}|Present|Current|Now|Ongoing)\b',
                r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\'?\d{2,4}\s*[-–—]\s*(?:Present|Current|Now|Ongoing)\b',
                r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
            ],
            'education': [r'\b(?:B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|M\.?B\.?A\.?|Ph\.?D\.?)\b', r'\b(?:University|College|Institute|School)\s+of\s+[A-Z][A-Za-z\s&]+\b', r'\b[A-Z][A-Za-z]+\s+(?:University|College|Institute)\b'],
            'companies': [r'\b[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*\s+(?:Inc|LLC|Ltd|Corporation|Corp|Co\.?)\b'],
            'contact': [r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,5}[-\s\.]?[0-9]{1,5}', r'(?i)linkedin\.com/in/[\w-]+', r'(?i)github\.com/[\w-]+'],
            'quantitative': [r'\$\d+(?:,\d{3})*(?:\.\d{2})?[KMB]?\b', r'\b\d+\+?\s*years?\b', r'\b\d+(?:\.\d+)?%\b', r'\b\d+x\b']
        }

    def _map_paragraph_to_section(self, para_index: int, sections: Dict[str, ResumeSection]) -> Optional[ResumeSection]:
        for section in sections.values():
            if para_index == section.title_index or para_index in section.paragraph_indices:
                return section
        return None

    def _identify_resume_sections(self, paragraphs: List[str]) -> Dict[str, ResumeSection]:
        sections = {}
        current_section_type = SectionType.UNKNOWN
        current_title = ""
        current_title_index = -1
        current_content = []
        current_indices = []
        def save_current_section():
            if current_title:
                sections[current_title] = ResumeSection(type=current_section_type, title=current_title, title_index=current_title_index, content=current_content, paragraph_indices=current_indices)
        for idx, para_text in enumerate(paragraphs):
            stripped_text = para_text.strip()
            if not stripped_text: continue
            section_type = self._identify_section_type(stripped_text)
            is_likely_header = (section_type != SectionType.UNKNOWN and len(stripped_text) < 50)
            if is_likely_header:
                save_current_section()
                current_section_type = section_type
                current_title = stripped_text
                current_title_index = idx
                current_content = []
                current_indices = []
            elif current_title:
                current_content.append(para_text)
                current_indices.append(idx)
        save_current_section()
        return sections

    def _identify_section_type(self, text: str) -> SectionType:
        for section_type, patterns in self.section_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    return section_type
        if "education" in text.lower() and "certification" in text.lower():
            return SectionType.EDUCATION 
        return SectionType.UNKNOWN

    def _sanitize_text(self, text: str) -> str:
        if not isinstance(text, str): 
            return ""
        return self.xml_illegal_char_re.sub('', text)

    def _identify_protected_content(self, text: str) -> List[Tuple[int, int]]:
        protected_regions = []
        for _, patterns in self.protection_patterns.items():
            for pattern in patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    protected_regions.append((match.start(), match.end()))
        protected_regions.sort()
        if not protected_regions: return []
        merged_regions = [protected_regions[0]]
        for current_start, current_end in protected_regions[1:]:
            last_start, last_end = merged_regions[-1]
            if current_start <= last_end:
                merged_regions[-1] = (last_start, max(last_end, current_end))
            else:
                merged_regions.append((current_start, current_end))
        return merged_regions

    def _create_protected_text_markup(self, text: str, protected_regions: List[Tuple[int, int]]) -> str:
        if not protected_regions: return text
        marked_text = ""
        last_end = 0
        for start, end in protected_regions:
            marked_text += text[last_end:start]
            marked_text += f" [PROTECTED: {text[start:end]}] "
            last_end = end
        marked_text += text[last_end:]
        return marked_text.strip()

    def _extract_enhanced_content(self, enhanced_text: str) -> str:
        return re.sub(r'\[PROTECTED:\s*([^\]]+)\]', r'\1', enhanced_text)

    def _construct_enhancement_system_prompt(self, sections_identified: Dict[str, ResumeSection]) -> str:
        sections_summary = ", ".join(f'"{title}"' for title in sections_identified.keys())
        return f"""You are an AI resume optimizer. Your task is to rewrite and enhance existing resume paragraphs to align with a job description.
DOCUMENT STRUCTURE: The resume has these sections: {sections_summary}.
YOUR TASKS:
1.  **Deep JD Alignment**: Rewrite EVERY paragraph provided to you to maximize its alignment with the job description.
2.  **Preserve Protected Content**: Text within [PROTECTED: ...] tags must remain untouched.
3.  **No New Content**: DO NOT generate any new paragraphs or sections. Your sole focus is on rewriting the content provided.
RESPONSE FORMAT:
Return a JSON object with a single key "enhanced_paragraphs". The value should be a dictionary where keys are the original paragraph IDs and values are the rewritten text."""

    def _construct_new_content_system_prompt(self, sections_identified: Dict[str, ResumeSection]) -> str:
        section_titles = [s.title.strip() for s in sections_identified.values()]
        sections_summary = ", ".join(f'"{title}"' for title in section_titles)
        return f"""You are an expert AI resume assistant. Your primary task is to generate new content for a resume based on user instructions.
HERE ARE THE EXACT SECTION TITLES AVAILABLE IN THE RESUME: {sections_summary}.
YOUR TASKS:
1.  **Interpret Instructions**: Based on the user's request, generate a complete and structured entry. For a new job experience, this includes the job title, company, responsibilities, etc.
2.  **Select Target Section**: Choose the best-fitting section title from the provided list.
3.  **Identify Anchor Text**: For adding to an EXISTING entry, this MUST be the **complete and unmodified text** of the paragraph you want to insert content AFTER. For a completely NEW entry, this MUST be `null`.
4.  **Generate Structured Content**: You must generate the content as an array of objects, where each object represents a line. This is critical for styling.
RESPONSE FORMAT:
Return a JSON object with a key "new_paragraphs". This must be an array of objects. Each object has:
- "target_section_title": An EXACT match to one of the section titles.
- "anchor_text": The unique anchor text for an update, or `null` for a NEW entry.
- "content": An array of **objects**. Each object must have:
    - "line_type": A style hint. Use 'job_header' for the main job title line, 'sub_header' for lines like 'Project:' or 'Responsibilities:', and 'bullet_point' for list items.
    - "text": The string content for the line.
Example for a NEW JOB entry:
{{"new_paragraphs": [{{
    "target_section_title": "Professional Experience",
    "anchor_text": null,
    "content": [
        {{ "line_type": "job_header", "text": "Principal Solutions Architect, ABC Techno Solutions LLC, Dallas, TX" }},
        {{ "line_type": "sub_header", "text": "Project: UiPath Enterprise Automation and Communication" }},
        {{ "line_type": "bullet_point", "text": "- Designed and maintained robust, enterprise-level architectures, ensuring alignment with organizational goals and technical requirements." }}
    ]
}}]}}"""

    def _construct_detailed_user_prompt(self, job_description: str, paragraphs_to_enhance: Dict) -> str:
        return f"""
        === JOB DESCRIPTION (PRIMARY OPTIMIZATION TARGET) ===
        {job_description}
        === RESUME CONTENT TO ENHANCE ===
        {json.dumps(paragraphs_to_enhance, indent=2)}
        === ENHANCEMENT REQUIREMENTS ===
        1.  Rewrite every paragraph to align deeply with the job description.
        2.  Preserve all [PROTECTED: ...] content exactly.
        3.  Return enhanced versions of ALL provided paragraphs. Do not add new ones."""

    def _construct_new_content_user_prompt(self, job_description: str, user_instructions: str) -> str:
        return f"""
        === USER INSTRUCTIONS ===
        {user_instructions}
        === JOB DESCRIPTION (FOR CONTEXT) ===
        {job_description}
        === TASK ===
        Based on the user instructions, generate the necessary new paragraphs. Your response must follow the structured format specified in your system prompt exactly."""

    def _apply_enhanced_text_with_format_preservation(self, paragraph, original_text, enhanced_text):
        try:
            sanitized_enhanced_text = self._sanitize_text(enhanced_text)
            if not sanitized_enhanced_text.strip():
                for run in paragraph.runs: paragraph._p.remove(run._r)
                return
            diffs = self.dmp.diff_main(original_text, sanitized_enhanced_text)
            self.dmp.diff_cleanupSemantic(diffs)
            original_runs_props = [{'text': r.text, 'bold': r.bold, 'italic': r.italic, 'underline': r.underline, 'font_name': r.font.name, 'font_size': r.font.size, 'font_color_rgb': r.font.color.rgb if r.font.color and r.font.color.rgb else None} for r in paragraph.runs]
            for run in paragraph.runs: paragraph._p.remove(run._r)
            if not original_runs_props:
                run = paragraph.add_run(sanitized_enhanced_text)
                run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
                return
            char_index_in_original = 0
            for op, data in diffs:
                if op == self.dmp.DIFF_DELETE:
                    char_index_in_original += len(data)
                    continue
                current_prop = None
                temp_char_index = 0
                for prop in original_runs_props:
                    prop_len = len(prop['text']) if prop['text'] else 0
                    if temp_char_index <= char_index_in_original < temp_char_index + prop_len:
                        current_prop = prop
                        break
                    temp_char_index += prop_len
                if not current_prop: current_prop = original_runs_props[-1]
                run = paragraph.add_run(data)
                if current_prop:
                    run.bold, run.italic, run.underline = current_prop.get('bold'), current_prop.get('italic'), current_prop.get('underline')
                    if current_prop.get('font_name'): run.font.name = current_prop['font_name']
                    if current_prop.get('font_size'): run.font.size = current_prop['font_size']
                    if current_prop.get('font_color_rgb'): run.font.color.rgb = current_prop['font_color_rgb']
                if op == self.dmp.DIFF_INSERT: run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
                if op == self.dmp.DIFF_EQUAL: char_index_in_original += len(data)
        except Exception as e:
            self.logger.error(f"Diff-match-patch failed: {e}. Falling back to simple replacement.")
            for run in paragraph.runs: paragraph._p.remove(run._r)
            new_run = paragraph.add_run(enhanced_text)
            new_run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    async def _async_enhance_chunk(self, session, paragraphs_chunk, job_description, system_prompt):
        user_prompt = self._construct_detailed_user_prompt(job_description, paragraphs_chunk)
        payload = {"model": "gpt-4o", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_object"}, "temperature": 0.3, "max_tokens": 16000}
        headers = {"Authorization": f"Bearer {os.getenv('OPENAI_API_KEY')}", "Content-Type": "application/json"}
        try:
            async with session.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers, timeout=aiohttp.ClientTimeout(total=300)) as response:
                if response.status != 200:
                    self.logger.error(f"API Error enhancing chunk. Status: {response.status}")
                    return paragraphs_chunk
                result = await response.json()
                content = json.loads(result["choices"][0]["message"]["content"])
                return content.get("enhanced_paragraphs", paragraphs_chunk)
        except Exception as e:
            self.logger.error(f"API call for enhancement chunk failed: {str(e)}")
            return paragraphs_chunk

    async def _async_generate_new_content(self, session, job_description, user_instructions, sections_map):
        if not user_instructions: return []
        system_prompt = self._construct_new_content_system_prompt(sections_map)
        user_prompt = self._construct_new_content_user_prompt(job_description, user_instructions)
        payload = {"model": "gpt-4o", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_object"}, "temperature": 0.5, "max_tokens": 16000}
        headers = {"Authorization": f"Bearer {os.getenv('OPENAI_API_KEY')}", "Content-Type": "application/json"}
        try:
            self.logger.info("Generating new content based on user instructions...")
            async with session.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers, timeout=aiohttp.ClientTimeout(total=180)) as response:
                if response.status != 200:
                    error_text = await response.text()
                    self.logger.error(f"API Error generating new content. Status: {response.status}, Details: {error_text}")
                    return []
                result = await response.json()
                content = json.loads(result["choices"][0]["message"]["content"])
                self.logger.info("Successfully generated new content.")
                return content.get("new_paragraphs", [])
        except Exception as e:
            self.logger.error(f"API call for new content failed: {str(e)}")
            return []

    def _get_all_paragraphs_in_order(self, doc: Document) -> List:
        paragraphs = []
        for child in doc._body._element:
            if child.tag == qn('w:p'):
                paragraphs.append(Paragraph(child, doc))
            elif child.tag == qn('w:tbl'):
                table = Table(child, doc)
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
        return paragraphs

    # --- THIS IS THE CORRECTED FUNCTION DEFINITION ---
    def enhance_resume(
        self,
        original_file_path: str,
        job_description: str,
        output_path: str,
        user_instructions: Optional[str] = None,
        progress_callback: Optional[Callable] = None # <-- ADDED THIS ARGUMENT
    ) -> dict:
    # --- END OF CORRECTION ---
        if not self.client: 
            return {"success": False, "error": "OpenAI client not initialized."}
        if not os.path.exists(original_file_path): 
            return {"success": False, "error": "Original file not found."}

        try:
            # --- Report Step 1 ---
            if progress_callback:
                progress_callback(step=1, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[0]["description"])

            doc = Document(original_file_path)
            all_paragraphs = self._get_all_paragraphs_in_order(doc)
            paragraph_map = {f"p_{i}": p for i, p in enumerate(all_paragraphs)}
            paragraph_texts = [p.text for p in all_paragraphs]
            sections_map = self._identify_resume_sections(paragraph_texts)
            self.logger.info(f"Identified section titles: {[title for title in sections_map.keys()]}")

            # --- Report Step 2 ---
            if progress_callback:
                progress_callback(step=2, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[1]["description"])

            paragraphs_for_ai = {}
            for i, p in enumerate(all_paragraphs):
                para_key = f"p_{i}"
                original_text = self._sanitize_text(p.text)
                if not original_text.strip(): 
                    continue
                
                protected_regions = self._identify_protected_content(original_text)
                paragraphs_for_ai[para_key] = self._create_protected_text_markup(original_text, protected_regions)

            if not paragraphs_for_ai and not user_instructions:
                doc.save(output_path)
                return {'success': True, 'summary': 'No content to enhance or add.'}

            async def run_enhancement_pipeline():
                async with aiohttp.ClientSession() as session:
                    
                    # --- Report Step 3 ---
                    if progress_callback:
                        progress_callback(step=3, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[2]["description"])

                    new_content_task = self._async_generate_new_content(session, job_description, user_instructions, sections_map)

                    enhancement_tasks = []
                    final_enhanced_paragraphs = {}

                    if paragraphs_for_ai:
                        # --- Report Step 4 ---
                        if progress_callback:
                            progress_callback(step=4, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[3]["description"])

                        CHUNK_SIZE = 40
                        paragraph_items = list(paragraphs_for_ai.items())
                        chunks = [dict(paragraph_items[i:i + CHUNK_SIZE]) for i in range(0, len(paragraph_items), CHUNK_SIZE)]
                        system_prompt = self._construct_enhancement_system_prompt(sections_map)

                        # --- Report Step 5 (Initialization) ---
                        chunk_count = len(chunks)
                        processed_chunks = 0
                        if progress_callback:
                            progress_callback(step=5, total_steps=self.TOTAL_STEPS, description=f"Enhancing resume content with AI (0/{chunk_count} chunks)")

                        async def _guarded(chunk):
                            nonlocal processed_chunks
                            result = await self._async_enhance_chunk(session, chunk, job_description, system_prompt)
                            processed_chunks += 1
                            # --- Report Step 5 (Progress) ---
                            if progress_callback:
                                progress_callback(step=5, total_steps=self.TOTAL_STEPS, description=f"Enhancing resume content with AI ({processed_chunks}/{chunk_count} chunks)")
                            return result

                        enhancement_tasks = [_guarded(chunk) for chunk in chunks]

                    # Gather with bounded concurrency
                    new_paragraph_entries, *enhanced_chunks = await asyncio.gather(new_content_task, *enhancement_tasks)
                    
                    for chunk_result in enhanced_chunks:
                        final_enhanced_paragraphs.update(chunk_result)
                    
                    # --- Report Step 5 (Completion) ---
                    if enhancement_tasks and progress_callback:
                         progress_callback(step=5, total_steps=self.TOTAL_STEPS, description=f"Enhancing resume content with AI ({processed_chunks}/{chunk_count} chunks complete)")

                    return {"enhanced": final_enhanced_paragraphs, "new": new_paragraph_entries or []}

            try:
                enhancement_result = asyncio.run(run_enhancement_pipeline())
            except Exception as api_error:
                self.logger.error(f"Enhancement pipeline failed: {str(api_error)}")
                return {"success": False, "error": f"API error: {str(api_error)}"}

            enhanced_marked_texts = enhancement_result["enhanced"]
            new_paragraph_entries = enhancement_result["new"]

            # --- Report Step 6 ---
            if progress_callback:
                progress_callback(step=6, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[5]["description"])

            enhancement_count = 0
            for key, enhanced_marked_text in enhanced_marked_texts.items():
                if key in paragraph_map:
                    original_para_text = self._sanitize_text(paragraph_map[key].text)
                    final_enhanced_text = self._extract_enhanced_content(enhanced_marked_text)
                    self._apply_enhanced_text_with_format_preservation(paragraph_map[key], original_para_text, final_enhanced_text)
                    if final_enhanced_text != original_para_text:
                        enhancement_count += 1
            
            # --- Report Step 7 ---
            if progress_callback:
                progress_callback(step=7, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[6]["description"])

            new_count = self._add_new_paragraphs_with_formatting(doc, new_paragraph_entries, sections_map, paragraph_map)

            self.logger.info(f"Applied {enhancement_count} paragraph enhancements.")
            self.logger.info(f"Added {new_count} new paragraphs.")

            # --- Report Step 8 ---
            if progress_callback:
                progress_callback(step=8, total_steps=self.TOTAL_STEPS, description=self.PROGRESS_STEPS[7]["description"])

            doc.save(output_path)

            return {
                "success": True,
                "instructions_applied": bool(user_instructions and new_count > 0),
                "enhancements_made": enhancement_count + new_count,
                "message": f"Successfully enhanced {enhancement_count} sections and added {new_count} new paragraphs."
            }

        except Exception as e:
            self.logger.exception(f"Enhancement pipeline failed: {str(e)}")
            return {"success": False, "error": str(e)}

    def _find_targeted_insertion_point(self, section_info, paragraph_map, anchor_text) -> Optional[Paragraph]:
        """ROBUST VERSION of the insertion point finder using a better fuzzy matching algorithm."""
        print("--- DIAGNOSTIC: _find_targeted_insertion_point (fuzzy v2) called ---")
        print(f" searching for anchor: '{anchor_text}'")
        
        best_match_score = -1
        best_match_para = None
        
        search_indices = [section_info.title_index] + section_info.paragraph_indices
        
        for idx in sorted(list(set(search_indices))):
            para = paragraph_map.get(f"p_{idx}")
            if para and para.text.strip():
                # token_set_ratio is excellent for matching phrases where words might be
                # out of order or some words are missing, like "Client:"
                score = fuzz.token_set_ratio(anchor_text, para.text)
                if score > best_match_score:
                    best_match_score = score
                    best_match_para = para

        # A lower threshold is fine for token_set_ratio as it's more accurate
        if best_match_score < 75:
            print(f"  [ERROR] Anchor text not found. Best match score was only {best_match_score} for '{anchor_text}'.")
            return None
            
        print(f"  Anchor found with score {best_match_score} in paragraph: '{best_match_para.text.strip()}'")
        print("--- DIAGNOSTIC: _find_targeted_insertion_point (fuzzy v2) finished ---")
        
        # We return the paragraph object of the best match, and the calling function will insert AFTER it.
        return best_match_para

    def _get_style_templates(self, section_info, paragraph_map):
        """DIAGNOSTIC VERSION of style template finder."""
        print("--- DIAGNOSTIC: _get_style_templates called ---")
        templates = {'job_header': None, 'sub_header': None, 'bullet_point': None}
        content_paras = [paragraph_map.get(f"p_{idx}") for idx in section_info.paragraph_indices if paragraph_map.get(f"p_{idx}")]
        if not content_paras:
            title_para = paragraph_map.get(f"p_{section_info.title_index}")
            for key in templates: templates[key] = title_para
            print("  Section has no content, using title paragraph as template for all.")
            return templates

        bullet_chars = ('✓', '•', '*', '-')
        for para in content_paras:
            if para.text.strip().startswith(bullet_chars) or 'list' in para.style.name.lower():
                templates['bullet_point'] = para
                print(f"  Found 'bullet_point' template: '{para.text.strip()}'")
                break
        for para in content_paras:
            if any(run.bold for run in para.runs):
                 templates['job_header'] = para
                 print(f"  Found 'job_header' template (bold text): '{para.text.strip()}'")
                 break
        for para in content_paras:
             if any(run.italic for run in para.runs) and not templates['job_header'] == para:
                 templates['sub_header'] = para
                 print(f"  Found 'sub_header' template (italic text): '{para.text.strip()}'")
                 break
        if not templates['job_header']:
            templates['job_header'] = content_paras[0]
            print(f"  No bold text found, using first content paragraph as 'job_header' template: '{content_paras[0].text.strip()}'")
        if not templates['sub_header']:
            templates['sub_header'] = templates['job_header']
            print("  No italic text found, using 'job_header' template for 'sub_header'.")
        if not templates['bullet_point']:
            templates['bullet_point'] = content_paras[-1]
            print(f"  No bullet character found, using last paragraph as 'bullet_point' template: '{content_paras[-1].text.strip()}'")

        print("--- DIAGNOSTIC: _get_style_templates finished ---")
        return templates

    def _add_new_paragraphs_with_formatting(self, doc, new_entries, sections_map, paragraph_map):
        """DIAGNOSTIC VERSION of the paragraph addition logic."""
        print("\n\n--- DIAGNOSTIC LOG START ---")
        if not new_entries:
            print("No new entries to add.")
            print("--- DIAGNOSTIC LOG END ---")
            return 0
        
        new_count = 0
        bullet_strip_pattern = re.compile(r'^[\s•✓*-]+\s*')
        for i, entry in enumerate(new_entries):
            print(f"\nProcessing new entry #{i+1}...")
            try:
                ai_section_title = entry.get("target_section_title")
                anchor_text = entry.get("anchor_text")
                content_items = entry.get("content", [])
                print(f"  AI wants to add to section: '{ai_section_title}'")
                print(f"  AI provided anchor text: '{anchor_text}'")
                print(f"  AI provided content items: {len(content_items)}")

                if not ai_section_title or not content_items:
                    print("  [SKIPPING] Entry is missing section title or content.")
                    continue

                found_section_info = None
                normalized_ai_title = ai_section_title.strip().lower().rstrip(':').strip()
                for title_key, section_info in sections_map.items():
                    normalized_doc_title = title_key.strip().lower().rstrip(':').strip()
                    if normalized_doc_title == normalized_ai_title:
                        found_section_info = section_info
                        print(f"  Successfully matched AI title '{ai_section_title}' to document section '{title_key}'")
                        break
                
                if not found_section_info:
                    print(f"  [ERROR] The section title '{ai_section_title}' was not found in the document's sections.")
                    continue
                
                insertion_point_para = None
                if anchor_text:
                    insertion_point_para = self._find_targeted_insertion_point(found_section_info, paragraph_map, anchor_text)
                else:
                    print("  No anchor text. Will insert at the top of the section.")
                    # Insert after the section title paragraph
                    insertion_point_para = paragraph_map.get(f"p_{found_section_info.title_index}")

                if not insertion_point_para:
                    print("  [ERROR] Could not determine a valid insertion point paragraph.")
                    continue
                
                style_templates = self._get_style_templates(found_section_info, paragraph_map)

                for item in reversed(content_items):
                    text_content = item.get('text', '')
                    line_type = item.get('line_type', 'bullet_point')
                    print(f"    - Preparing to insert line: '{text_content}' with style '{line_type}'")
                    template_para = style_templates.get(line_type)
                    if not template_para:
                        template_para = style_templates['bullet_point']
                    
                    new_p_element = copy.deepcopy(template_para._p)
                    new_para = Paragraph(new_p_element, doc)
                    for run in new_para.runs:
                        p = new_para._p
                        p.remove(run._r)
                    
                    clean_text = bullet_strip_pattern.sub('', text_content) if line_type == 'bullet_point' else text_content
                    new_run = new_para.add_run(clean_text)
                    new_run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
                    
                    if line_type == 'job_header' and style_templates.get('job_header') and any(r.bold for r in style_templates['job_header'].runs):
                        new_run.bold = True
                    
                    insertion_point_para._p.addnext(new_p_element)
                    new_count += 1
                print("  Successfully inserted all content items for this entry.")

            except Exception as e:
                self.logger.exception(f"Critical error processing new entry: {e}")
                print(f"  [CRITICAL ERROR] An exception occurred: {e}")
        
        print("\n--- DIAGNOSTIC LOG END ---\n")
        return new_count